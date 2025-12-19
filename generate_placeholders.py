import os
from docx import Document
from fpdf import FPDF
from openpyxl import Workbook

# Define the documents to generate
documents = [
    {"filename": "Change_Order_Proposal_Template.pdf", "title": "Change Order Proposal Template", "type": "pdf"},
    {"filename": "Gemini_FY25_Pricing_Guide.xlsx", "title": "Gemini FY25 Pricing Guide", "type": "xlsx"},
    {"filename": "Closeout_Maintenance_Warranty.docx", "title": "Closeout: Maintenance & Warranty", "type": "docx"},
    {"filename": "Signage_Takeoff_Example.docx", "title": "Signage Takeoff Example", "type": "docx"},
    {"filename": "Warranty_Maintenance_Info.docx", "title": "Warranty & Maintenance Info", "type": "docx"},
    {"filename": "Unconditional_Sub_Lien_Waiver.docx", "title": "Unconditional Sub Lien Waiver", "type": "docx"},
    {"filename": "Change_Order_Example_LCS.pdf", "title": "Change Order Example (LCS)", "type": "pdf"},
    {"filename": "Sample_COI_Moon_River.pdf", "title": "Sample COI for Moon River", "type": "pdf"},
    {"filename": "Combined_Lien_Waiver_Sample.pdf", "title": "Combined Lien Waiver Sample", "type": "pdf"},
    {"filename": "COI_Requirements_Template.docx", "title": "COI Requirements Template", "type": "docx"},
    {"filename": "Schedule_of_Values_Example.docx", "title": "Schedule of Values Example", "type": "docx"},
    {"filename": "Trello_Project_Setup_Checklist.docx", "title": "Trello Project Setup Checklist", "type": "docx"},
    {"filename": "Latest_Bid_Document_Template.docx", "title": "Latest Bid Document Template", "type": "docx"},
    {"filename": "RFP_Response_Example.docx", "title": "RFP Response Example", "type": "docx"}
]

output_dir = "/home/ubuntu/moonriver-knowledge-base/public/documents"
os.makedirs(output_dir, exist_ok=True)

def create_pdf(filename, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=16)
    pdf.cell(200, 10, txt=title, ln=1, align='C')
    pdf.set_font("Arial", size=12)
    pdf.ln(20)
    pdf.multi_cell(0, 10, txt=f"This is a placeholder document for '{title}'.\n\nReplace this file with the actual content when available.")
    pdf.output(os.path.join(output_dir, filename))
    print(f"Created PDF: {filename}")

def create_docx(filename, title):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"This is a placeholder document for '{title}'.")
    doc.add_paragraph("Replace this file with the actual content when available.")
    doc.save(os.path.join(output_dir, filename))
    print(f"Created DOCX: {filename}")

def create_xlsx(filename, title):
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws['A1'] = title
    ws['A3'] = "This is a placeholder spreadsheet."
    ws['A4'] = "Replace this file with the actual content when available."
    wb.save(os.path.join(output_dir, filename))
    print(f"Created XLSX: {filename}")

for doc in documents:
    if doc["type"] == "pdf":
        create_pdf(doc["filename"], doc["title"])
    elif doc["type"] == "docx":
        create_docx(doc["filename"], doc["title"])
    elif doc["type"] == "xlsx":
        create_xlsx(doc["filename"], doc["title"])
